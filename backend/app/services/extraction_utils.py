"""
Motor de extracción estructurada de documentos para el módulo OC y Financiero.

Extrae pares {campo_canónico: valor_str} desde documentos Excel y Word usando
el diccionario de sinónimos de field_synonyms.py.

Punto de verdad único — importar desde aquí en cotizaciones.py y facturas.py.
"""

import io
import logging
from typing import Optional

from app.services.field_synonyms import fuzzy_resolve, resolve_field

log = logging.getLogger(__name__)


def extraer_campos_estructurado(contenido: bytes, ext: str) -> dict[str, str]:
    """Extrae pares etiqueta→valor de documentos estructurados (Excel, Word).

    Usa field_synonyms para normalizar los encabezados a nombres canónicos.
    Retorna un dict {campo_canonico: valor_str} con los campos reconocidos.
    Retorna dict vacío si el formato no es Excel/Word o si falla la extracción.
    """
    resultado: dict[str, str] = {}

    if ext in ("xlsx", "xls"):
        try:
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(contenido), data_only=True)
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    cells = [c for c in row if c is not None]
                    if len(cells) < 2:
                        continue
                    label = str(cells[0]).strip()
                    value = str(cells[1]).strip()
                    if not label or not value or value.lower() == "none":
                        continue
                    canonical = resolve_field(label) or (fuzzy_resolve(label)[0])
                    if canonical and canonical not in resultado:
                        resultado[canonical] = value
        except Exception as e:
            log.warning("[extraccion] Excel estructurado falló: %s", e)

    elif ext == "docx":
        try:
            from docx import Document

            doc = Document(io.BytesIO(contenido))

            # Tablas: primera columna = etiqueta, segunda columna = valor
            for table in doc.tables:
                for row in table.rows:
                    raw_cells = [c.text.strip() for c in row.cells]
                    # Deduplicar celdas combinadas (python-docx las repite)
                    deduped: list[str] = []
                    for c in raw_cells:
                        if not deduped or c != deduped[-1]:
                            deduped.append(c)
                    cells = [c for c in deduped if c]
                    if len(cells) < 2:
                        continue
                    label, value = cells[0], cells[1]
                    canonical = resolve_field(label) or (fuzzy_resolve(label)[0])
                    if canonical and canonical not in resultado:
                        resultado[canonical] = value

            # Párrafos con formato "Etiqueta: valor"
            for para in doc.paragraphs:
                text = para.text.strip()
                if ":" not in text:
                    continue
                label, _, value = text.partition(":")
                label = label.strip()
                value = value.strip()
                if not value:
                    continue
                canonical = resolve_field(label) or (fuzzy_resolve(label)[0])
                if canonical and canonical not in resultado:
                    resultado[canonical] = value

        except Exception as e:
            log.warning("[extraccion] Word estructurado falló: %s", e)

    return resultado
