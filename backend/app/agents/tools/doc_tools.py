"""
Herramientas de documentos RAG para los agentes ZYMO.

Estas funciones son las que Gemini llama como "tools" cuando necesita
buscar en el conocimiento interno de ZYMO.
"""
import io
import logging
from pathlib import Path
from typing import Optional

from sqlmodel import Session, select

from app.agent_database import AgentDocumento, get_agents_engine
from app.agents.lightrag_service import buscar_conocimiento, indexar_texto
from app.config import settings

logger = logging.getLogger(__name__)


# ── Extracción de texto plano ──────────────────────────────────────────────────

def extraer_texto(contenido: bytes, extension: str) -> str:
    """
    Extrae texto plano de un archivo.
    Soporta: .md, .txt, .pdf, .docx
    """
    ext = extension.lower().lstrip(".")

    if ext in ("md", "txt"):
        return contenido.decode("utf-8", errors="ignore")

    if ext == "pdf":
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(contenido)) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages)
        except Exception as e:
            logger.warning("Error extrayendo PDF: %s", e)
            return ""

    if ext == "docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(contenido))
            partes = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    partes.append(" | ".join(c.text.strip() for c in row.cells if c.text.strip()))
            return "\n".join(partes)
        except Exception as e:
            logger.warning("Error extrayendo DOCX: %s", e)
            return ""

    logger.warning("Formato no soportado para RAG: %s", ext)
    return ""


# ── Operaciones sobre documentos ──────────────────────────────────────────────

async def subir_e_indexar(
    nombre: str,
    extension: str,
    contenido: bytes,
    tipo: str,
    area: str,
    subido_por_id: int,
) -> AgentDocumento:
    """
    Extrae texto del archivo, lo indexa en LightRAG y guarda metadata en BD.
    Retorna el registro AgentDocumento creado.
    """
    texto = extraer_texto(contenido, extension)

    # Guardar el archivo físico en agent_docs/
    docs_dir = Path(settings.agent_docs_dir) / area
    docs_dir.mkdir(parents=True, exist_ok=True)
    ruta = str(docs_dir / f"{nombre}.{extension.lstrip('.')}")
    Path(ruta).write_bytes(contenido)

    # Indexar en LightRAG (si falla, igual guardamos la metadata)
    indexado = False
    if texto:
        indexado = await indexar_texto(texto)

    chunks_estimados = len(texto) // 500 if indexado else 0

    doc = AgentDocumento(
        nombre=nombre,
        ruta=ruta,
        tipo=tipo,
        area=area,
        chunks_count=chunks_estimados,
        subido_por_id=subido_por_id,
    )
    with Session(get_agents_engine()) as db:
        db.add(doc)
        db.commit()
        db.refresh(doc)

    logger.info(
        "Documento '%s' indexado: %d chars, %d chunks estimados, indexado=%s",
        nombre, len(texto), chunks_estimados, indexado,
    )
    return doc


def listar_documentos(area: Optional[str] = None) -> list[AgentDocumento]:
    """Lista documentos registrados, opcionalmente filtrados por área."""
    with Session(get_agents_engine()) as db:
        query = select(AgentDocumento)
        if area:
            query = query.where(AgentDocumento.area == area)
        return list(db.exec(query).all())


def obtener_documento(doc_id: str) -> Optional[AgentDocumento]:
    with Session(get_agents_engine()) as db:
        return db.get(AgentDocumento, doc_id)


def eliminar_documento(doc_id: str) -> bool:
    """Elimina el registro de BD y el archivo físico. El grafo LightRAG no se modifica (limitación conocida)."""
    with Session(get_agents_engine()) as db:
        doc = db.get(AgentDocumento, doc_id)
        if not doc:
            return False
        # Intentar borrar el archivo físico
        try:
            Path(doc.ruta).unlink(missing_ok=True)
        except Exception as e:
            logger.warning("No se pudo borrar archivo %s: %s", doc.ruta, e)
        db.delete(doc)
        db.commit()
    return True


# ── Tool que llama Gemini ──────────────────────────────────────────────────────

async def buscar_en_conocimiento(
    query: str,
    area: Optional[str] = None,
    modo: str = "mix",
) -> dict:
    """
    Tool principal: busca en el grafo de conocimiento ZYMO.

    Args:
        query: Pregunta o término a buscar
        area: Filtro informativo (LightRAG busca en todo el grafo)
        modo: "local" | "global" | "mix"

    Returns:
        dict con resultado, query original y metadata
    """
    resultado = await buscar_conocimiento(query, modo=modo)
    return {
        "query": query,
        "resultado": resultado,
        "area_filtro": area,
        "modo": modo,
    }
