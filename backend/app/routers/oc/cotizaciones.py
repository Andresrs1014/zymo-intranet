import io
import json
import logging
import threading
import re
import uuid
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Optional

log = logging.getLogger(__name__)

from fastapi import APIRouter, BackgroundTasks, Depends, File, HTTPException, UploadFile, status
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from sqlmodel import Session, select

from app.core.deps import get_current_user, require_compras
from app.core.permissions import user_has_permission
from app.database import get_db
from app.oc_database import get_oc_db
from app.models.oc import CotizacionProveedor, EstadoOC, SolicitudOC
from app.models.user import User
from app.routers.oc.documentos import regenerar_pdf_orden_por_solicitud
from app.services.email_service import (
    send_aprobacion_directora,
    send_alerta_correccion_post_cierre,
    send_correccion_directivo,
    send_cotizacion_lista,
    send_oc_a_proveedor,
    send_proforma_financiero,
)
from app.services.historial import registrar_cambio_estado
from app.config import settings

router = APIRouter(tags=["OC - Cotizaciones"])


# ── Auto-discovery de campos no reconocidos ───────────────────────────────────

_CANDIDATES_FILE = Path("/app/data/field_candidates.json")
_candidates_lock = threading.Lock()


def _registrar_candidatos_campo(candidatos: list[str], fuente: str) -> None:
    """Registra etiquetas no reconocidas como candidatos para nuevos sinónimos.

    Args:
        candidatos: Lista de etiquetas no reconocidas encontradas en el documento.
        fuente: Descripción de dónde se encontró (ej: "pdf_tabla", "excel", "docx").
    """
    if not candidatos:
        return
    try:
        with _candidates_lock:
            _CANDIDATES_FILE.parent.mkdir(parents=True, exist_ok=True)
            existing: list[dict] = []
            if _CANDIDATES_FILE.exists():
                try:
                    existing = json.loads(_CANDIDATES_FILE.read_text(encoding="utf-8"))
                except (json.JSONDecodeError, OSError):
                    existing = []

            hoy = date.today().isoformat()
            labels_existentes = {
                e.get("label", "").lower() for e in existing if e.get("label")
            }
            nuevos = [
                {"label": c.strip().lower(), "fuente": fuente, "fecha": hoy}
                for c in candidatos
                if c.strip().lower() not in labels_existentes and len(c.strip()) > 1
            ]
            if nuevos:
                existing.extend(nuevos)
                _CANDIDATES_FILE.write_text(
                    json.dumps(existing, ensure_ascii=False, indent=2),
                    encoding="utf-8",
                )
                log.info("[auto-discovery] %d nuevos candidatos de campo registrados desde %s: %s",
                         len(nuevos), fuente, [n["label"] for n in nuevos])
    except Exception as e:
        log.warning("[auto-discovery] No se pudieron registrar candidatos: %s", e)


# ── Schemas ───────────────────────────────────────────────────────────────────

class CotizacionCreate(BaseModel):
    proveedor_nombre: str
    proveedor_nit: Optional[str] = None
    proveedor_email: Optional[str] = None
    numero_cotizacion_proveedor: Optional[str] = None
    valor_unitario: float = 0
    valor_antes_iva: Optional[float] = None
    valor_iva: Optional[float] = None
    valor_total: float
    fecha_estimada_entrega: Optional[date] = None
    forma_pago: Optional[str] = None
    plazo_entrega: Optional[str] = None
    garantia: Optional[str] = None
    anticipo: Optional[str] = None
    pago_saldo: Optional[str] = None
    observaciones: Optional[str] = None
    # Lista de ítems; None o vacía = cotización de un solo producto
    items: Optional[list[dict]] = None


class CotizacionRead(BaseModel):
    id: uuid.UUID
    solicitud_id: uuid.UUID
    proveedor_nombre: str
    proveedor_nit: Optional[str]
    proveedor_email: Optional[str]
    numero_cotizacion_proveedor: Optional[str]
    valor_unitario: float
    valor_antes_iva: Optional[float]
    valor_iva: Optional[float]
    valor_total: float
    fecha_estimada_entrega: Optional[date]
    forma_pago: Optional[str]
    plazo_entrega: Optional[str]
    garantia: Optional[str]
    anticipo: Optional[str]
    pago_saldo: Optional[str]
    observaciones: Optional[str]
    items: Optional[list] = None
    pdf_path: Optional[str]
    extraccion_automatica: bool
    aprobada: Optional[bool]
    valor_aprobado: Optional[float]
    valor_aprobado_original: Optional[float] = None
    aprobado_por_id: Optional[int]
    observaciones_aprobacion: Optional[str]
    created_at: datetime

    class Config:
        from_attributes = True


class AprobarPayload(BaseModel):
    valor_aprobado: float
    observaciones_aprobacion: Optional[str] = None


class RechazarPayload(BaseModel):
    observaciones_aprobacion: str


class CancelarCotizacionPayload(BaseModel):
    justificacion: str


class CorreccionCotizacionPayload(BaseModel):
    que_corregir: str
    destino: str  # "auxiliar" | "solicitante"


class EditarCotizacionPayload(BaseModel):
    proveedor_nombre: Optional[str] = None
    proveedor_nit: Optional[str] = None
    proveedor_email: Optional[str] = None
    numero_cotizacion_proveedor: Optional[str] = None
    valor_antes_iva: Optional[float] = None
    valor_iva: Optional[float] = None
    valor_total: Optional[float] = None
    forma_pago: Optional[str] = None
    plazo_entrega: Optional[str] = None
    observaciones: Optional[str] = None
    items: Optional[list[dict]] = None


class CorregirDirectivoPayload(BaseModel):
    # Campos editables de la cotización (todos opcionales)
    proveedor_nombre: Optional[str] = None
    proveedor_nit: Optional[str] = None
    proveedor_email: Optional[str] = None
    numero_cotizacion_proveedor: Optional[str] = None
    valor_antes_iva: Optional[float] = None
    valor_iva: Optional[float] = None
    valor_total: Optional[float] = None
    forma_pago: Optional[str] = None
    plazo_entrega: Optional[str] = None
    observaciones: Optional[str] = None
    items: Optional[list[dict]] = None
    # Puede cambiar el valor aprobado
    valor_aprobado: Optional[float] = None
    # Nota opcional del director sobre la corrección
    observacion_correccion: Optional[str] = Field(default=None, max_length=1000)
    # Justificación requerida cuando el proceso ya está cerrado/entregado
    motivo_post_cierre: Optional[str] = Field(default=None, max_length=2000)


class ItemCotizacion(BaseModel):
    num: Optional[int] = None
    descripcion: str
    referencia: Optional[str] = None
    cantidad: Optional[float] = None
    valor_unitario: Optional[float] = None
    valor_total: Optional[float] = None


class ExtraccionResult(BaseModel):
    # Campos escalares del encabezado de la cotización
    proveedor_nit: Optional[str] = None
    numero_cotizacion_proveedor: Optional[str] = None
    proveedor_nombre: Optional[str] = None
    valor_unitario: Optional[float] = None
    valor_antes_iva: Optional[float] = None
    valor_iva: Optional[float] = None
    valor_total: Optional[float] = None
    forma_pago: Optional[str] = None
    plazo_entrega: Optional[str] = None
    garantia: Optional[str] = None
    anticipo: Optional[str] = None
    pago_saldo: Optional[str] = None
    # Tabla de ítems (vacía si el documento tiene un solo producto)
    items: list[ItemCotizacion] = []
    nombre_archivo: str = ""
    campos_encontrados: int = 0
    # Referencia al archivo temporal guardado — se vincula al crear la cotización
    temp_file_ext: Optional[str] = None


# ── Helpers de extracción ─────────────────────────────────────────────────────

def _extraer_texto(contenido: bytes, ext: str) -> str:
    if ext == "pdf":
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(contenido)) as pdf:
                texto = "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception as e:
            log.warning("[motor] extracción texto PDF falló: %s", e)
            texto = ""
        from app.services.ocr_service import texto_con_ocr_fallback
        return texto_con_ocr_fallback(texto, contenido)
    if ext in ("xlsx", "xls"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(contenido), data_only=True)
            lines: list[str] = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    parts = [str(c) for c in row if c is not None]
                    if parts:
                        lines.append("  ".join(parts))
            return "\n".join(lines)
        except Exception as e:
            log.warning("[motor] extracción texto Excel falló: %s", e)
            return ""
    if ext == "docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(contenido))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            log.warning("[motor] extracción texto Word falló: %s", e)
            return ""
    return ""


from app.services.extraction_utils import extraer_campos_estructurado as _extraer_campos_estructurado  # noqa: E402


from app.services.number_utils import parse_cop as _to_float  # noqa: E402
from app.services.cotizacion_parse import parsear_campos_cotizacion  # noqa: E402


def _parsear_campos(texto: str, extra: Optional[dict[str, str]] = None) -> ExtraccionResult:
    data = parsear_campos_cotizacion(texto, extra)
    values = [
        data.get("proveedor_nit"),
        data.get("numero_cotizacion_proveedor"),
        data.get("proveedor_nombre"),
        data.get("valor_total"),
        data.get("valor_antes_iva"),
        data.get("valor_iva"),
        data.get("valor_unitario"),
        data.get("forma_pago"),
        data.get("plazo_entrega"),
        data.get("garantia"),
        data.get("anticipo"),
        data.get("pago_saldo"),
    ]
    campos = sum(1 for v in values if v is not None)

    return ExtraccionResult(
        proveedor_nit=data.get("proveedor_nit"),
        numero_cotizacion_proveedor=data.get("numero_cotizacion_proveedor"),
        proveedor_nombre=data.get("proveedor_nombre"),
        valor_unitario=data.get("valor_unitario"),
        valor_antes_iva=data.get("valor_antes_iva"),
        valor_iva=data.get("valor_iva"),
        valor_total=data.get("valor_total"),
        forma_pago=data.get("forma_pago"),
        plazo_entrega=data.get("plazo_entrega"),
        garantia=data.get("garantia"),
        anticipo=data.get("anticipo"),
        pago_saldo=data.get("pago_saldo"),
        campos_encontrados=campos,
    )


# ── Motor de extracción de tabla de ítems ─────────────────────────────────────
#
# Detecta la tabla de productos en cualquier cotización buscando una fila de
# encabezado cuyos títulos resuelvan a campos canónicos conocidos (descripcion,
# cantidad, valor_unitario, valor_total, referencia) usando field_synonyms.
#
# Soporta: Excel (.xlsx/.xls), PDF (via pdfplumber.extract_tables), Word (.docx)
#
# El motor regex existente (_parsear_campos) sigue activo para los campos
# escalares del encabezado (NIT, forma de pago, totales, etc.).

_CAMPOS_TABLA = {"descripcion", "cantidad", "valor_unitario", "valor_total", "referencia"}


def _fila_a_item(col_map: dict[int, str], celda_valores: list[str]) -> Optional[dict]:
    """Convierte una fila de datos a un dict de ítem usando el mapa de columnas.

    Retorna None si la fila no tiene descripción (fila vacía o de totales).
    """
    item: dict = {}
    for col_idx, canonical in col_map.items():
        if col_idx >= len(celda_valores):
            continue
        val = celda_valores[col_idx].strip()
        if not val or val.lower() in ("none", "n/a", "-", "—"):
            continue
        # Convertir campos numéricos
        if canonical in ("cantidad", "valor_unitario", "valor_total"):
            parsed = _to_float(val)
            item[canonical] = parsed if parsed is not None else val
        else:
            item[canonical] = val

    return item if item.get("descripcion") else None


def _detectar_encabezado(filas_celdas: list[list[str]]) -> Optional[tuple[int, dict[int, str]]]:
    """Busca la primera fila que sea un encabezado de tabla de ítems.

    Criterio: la fila debe resolver al menos 'descripcion' + uno de
    ('cantidad', 'valor_unitario', 'valor_total') usando field_synonyms.

    Retorna (índice_fila, {col_idx: campo_canónico}) o None.
    """
    from app.services.field_synonyms import resolve_field, fuzzy_resolve

    for row_idx, celdas in enumerate(filas_celdas):
        col_map: dict[int, str] = {}
        for col_idx, celda in enumerate(celdas):
            if not celda:
                continue
            canonical = resolve_field(celda)
            if not canonical:
                canonical, score = fuzzy_resolve(celda, threshold=0.72)
            if canonical and canonical in _CAMPOS_TABLA:
                col_map[col_idx] = canonical

        tiene_desc = "descripcion" in col_map.values()
        tiene_valor = any(f in col_map.values() for f in ("cantidad", "valor_unitario", "valor_total"))

        if tiene_desc and tiene_valor:
            # Recolectar etiquetas no reconocidas de esta fila de encabezado
            no_reconocidas = [
                celda.strip()
                for i, celda in enumerate(celdas)
                if celda.strip() and i not in col_map
            ]
            if no_reconocidas:
                _registrar_candidatos_campo(no_reconocidas, "encabezado_tabla")
            return row_idx, col_map

    return None


def _items_desde_excel(contenido: bytes) -> list[dict]:
    """Extrae tabla de ítems de un Excel buscando la fila de encabezado."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(contenido), data_only=True)
        for ws in wb.worksheets:
            filas = [
                [str(c).strip() if c is not None else "" for c in row]
                for row in ws.iter_rows(values_only=True)
            ]
            resultado = _detectar_encabezado(filas)
            if resultado is None:
                continue
            header_idx, col_map = resultado
            items = []
            for fila in filas[header_idx + 1:]:
                if not any(fila):
                    continue
                item = _fila_a_item(col_map, fila)
                if item:
                    items.append(item)
            if items:
                return items
    except Exception as e:
        log.warning("[motor] extracción ítems Excel falló: %s", e)
    return []


def _items_desde_pdf_tablas(contenido: bytes) -> list[dict]:
    """Extrae ítems de un PDF con tablas con bordes visibles (pdfplumber.extract_tables)."""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(contenido)) as pdf:
            for page in pdf.pages:
                for tabla in (page.extract_tables() or []):
                    if not tabla or len(tabla) < 2:
                        continue
                    filas = [
                        [str(c).strip() if c else "" for c in fila]
                        for fila in tabla
                    ]
                    resultado = _detectar_encabezado(filas)
                    if resultado is None:
                        continue
                    header_idx, col_map = resultado
                    items = []
                    for fila in filas[header_idx + 1:]:
                        if not any(fila):
                            continue
                        item = _fila_a_item(col_map, fila)
                        if item:
                            items.append(item)
                    if items:
                        return items
    except Exception as e:
        log.warning("[motor] extracción ítems PDF (tablas) falló: %s", e)
    return []


def _fusionar_palabras_por_gap(
    palabras: list[dict], gap_threshold: float = 1.8
) -> list[str]:
    """Fusiona palabras consecutivas de una línea cuando el espacio horizontal
    entre ellas (next.x0 - prev.x1) es menor a ``gap_threshold`` puntos.

    Esto evita pegar columnas distintas (p. ej. ``25.090`` y ``100.360``) que
    visualmente están separadas, y al mismo tiempo une fragmentos del mismo
    número que pdfplumber reportó por separado por culpa del kerning.
    """
    resultado: list[str] = []
    if not palabras:
        return resultado

    cur = str(palabras[0]["text"])
    cur_x1 = float(palabras[0]["x1"])
    for w in palabras[1:]:
        x0 = float(w["x0"])
        gap = x0 - cur_x1
        if gap < gap_threshold:
            cur += str(w["text"])
        else:
            resultado.append(cur)
            cur = str(w["text"])
        cur_x1 = float(w["x1"])
    resultado.append(cur)
    return resultado


def _items_desde_pdf_texto(contenido: bytes) -> list[dict]:
    """Fallback: extrae ítems de PDFs sin bordes de tabla, usando posición X/Y de palabras.

    Agrupa palabras por línea (mismo Y aproximado) y detecta líneas con patrón
    descripción + valor(es) numérico(s) al final — estructura típica de PDFs
    de proveedores colombianos generados sin tablas explícitas.
    """
    _ENCABEZADOS_SKIP = {
        "descripción", "descripcion", "total", "subtotal", "iva",
        "cantidad", "cant", "valor", "precio", "item", "ítem",
        "referencia", "ref", "und", "unidad",
    }
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(contenido)) as pdf:
            for page in pdf.pages:
                words = page.extract_words(x_tolerance=3, y_tolerance=3)
                if not words:
                    continue

                words_sorted = sorted(words, key=lambda w: (float(w["top"]), float(w["x0"])))
                lineas_words: list[list[dict]] = []
                for w in words_sorted:
                    t = float(w["top"])
                    if not lineas_words:
                        lineas_words.append([w])
                        continue
                    ref_top = float(lineas_words[-1][0]["top"])
                    if abs(ref_top - t) <= 3.0:
                        lineas_words[-1].append(w)
                    else:
                        lineas_words.append([w])

                items: list[dict] = []
                for linea in lineas_words:
                    linea.sort(key=lambda w: float(w["x0"]))
                    tokens = _fusionar_palabras_por_gap(linea)
                    if len(tokens) < 2:
                        continue

                    # Separar tokens numéricos al final de la línea
                    resto = list(tokens)
                    valores_raw: list[str] = []
                    while resto and re.match(r"^[\d.,\$%]+$", resto[-1].replace(".", "").replace(",", "")):
                        valores_raw.insert(0, resto.pop())

                    if not valores_raw or not resto:
                        continue

                    # Saltar líneas que sean encabezados de tabla o totales
                    primer_token = resto[0].lower().rstrip(":")
                    if primer_token in _ENCABEZADOS_SKIP:
                        continue
                    if any(t.upper() in ("TOTAL", "SUBTOTAL", "IVA", "GRAN") for t in resto):
                        continue

                    cantidad: Optional[float] = None
                    if (
                        len(resto) >= 2
                        and re.fullmatch(r"\d{1,4}", resto[0])
                        and len(valores_raw) >= 1
                    ):
                        cantidad = float(resto[0])
                        resto = resto[1:]

                    descripcion = " ".join(resto)
                    valores = [v for v in (_to_float(n) for n in valores_raw) if v is not None]

                    item: dict = {"descripcion": descripcion}
                    if cantidad is not None:
                        item["cantidad"] = cantidad
                    if len(valores) >= 2:
                        item["valor_unitario"] = valores[-2]
                        item["valor_total"] = valores[-1]
                    elif len(valores) == 1:
                        if cantidad is not None and cantidad > 0:
                            item["valor_unitario"] = valores[0]
                            item["valor_total"] = valores[0] * cantidad
                        else:
                            item["valor_total"] = valores[0]

                    items.append(item)

                if len(items) >= 1:
                    log.info("[motor] PDF texto: %d ítems detectados por posición", len(items))
                    return items
    except Exception as e:
        log.warning("[motor] extracción ítems PDF (texto) falló: %s", e)
    return []


def _items_desde_pdf(contenido: bytes) -> list[dict]:
    """Extrae tabla de ítems de un PDF.

    Intenta primero con tablas explícitas (bordes visibles). Si no encuentra
    nada, usa el fallback por posición de texto para PDFs sin bordes.
    """
    items = _items_desde_pdf_tablas(contenido)
    if items:
        return items
    log.info("[motor] PDF sin tablas detectables — usando fallback por posición de texto")
    return _items_desde_pdf_texto(contenido)


def _items_desde_docx(contenido: bytes) -> list[dict]:
    """Extrae tabla de ítems de un Word buscando la tabla con encabezado de productos."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(contenido))
        for tabla in doc.tables:
            if len(tabla.rows) < 2:
                continue
            # Convertir cada fila a lista de strings, deduplicando celdas combinadas
            filas: list[list[str]] = []
            for row in tabla.rows:
                celdas_raw = [c.text.strip() for c in row.cells]
                deduped: list[str] = []
                for c in celdas_raw:
                    if not deduped or c != deduped[-1]:
                        deduped.append(c)
                filas.append(deduped)

            resultado = _detectar_encabezado(filas)
            if resultado is None:
                continue
            header_idx, col_map = resultado
            items = []
            for fila in filas[header_idx + 1:]:
                if not any(fila):
                    continue
                item = _fila_a_item(col_map, fila)
                if item:
                    items.append(item)
            if items:
                return items
    except Exception as e:
        log.warning("[motor] extracción ítems Word falló: %s", e)
    return []


def _extraer_tabla_items(contenido: bytes, ext: str) -> list[dict]:
    """Punto de entrada del motor de extracción de ítems.

    Intenta detectar y extraer la tabla de productos de una cotización.
    Retorna lista vacía si el documento tiene formato de un solo producto.
    """
    if ext in ("xlsx", "xls"):
        return _items_desde_excel(contenido)
    if ext == "pdf":
        return _items_desde_pdf(contenido)
    if ext == "docx":
        return _items_desde_docx(contenido)
    return []


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.post(
    "/solicitudes/{solicitud_id}/cotizacion/extraer",
    response_model=ExtraccionResult,
    status_code=status.HTTP_200_OK,
)
async def extraer_cotizacion(
    solicitud_id: uuid.UUID,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    current_user: User = Depends(require_compras),
    oc_db: Session = Depends(get_oc_db),
):
    """Extrae campos de un PDF/Excel/Word de cotización sin guardar nada.

    Fase 1 (síncrona): regex + sinónimos → respuesta inmediata.
    Fase 2 (background): Gemini completa campos vacíos, alimenta aprendizaje.
    El frontend hace poll a GET .../cotizacion/extraccion/resultado para Fase 2.
    """
    from app.services.extraction_pipeline import run_phase1, run_phase2

    solicitud = oc_db.get(SolicitudOC, solicitud_id)
    if not solicitud:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Solicitud no encontrada.")

    nombre = file.filename or ""
    ext = nombre.rsplit(".", 1)[-1].lower() if "." in nombre else ""
    if ext not in ("pdf", "xlsx", "xls", "docx"):
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Formato no soportado. Use PDF, Excel (.xlsx) o Word (.docx).",
        )

    contenido = await file.read()

    # ── Fase 1: síncrona ──────────────────────────────────────────────────────
    phase1_dict = run_phase1(contenido, ext, str(solicitud_id))

    resultado = ExtraccionResult(
        proveedor_nit=phase1_dict.get("proveedor_nit"),
        numero_cotizacion_proveedor=phase1_dict.get("numero_cotizacion_proveedor"),
        proveedor_nombre=phase1_dict.get("proveedor_nombre"),
        valor_unitario=phase1_dict.get("valor_unitario"),
        valor_antes_iva=phase1_dict.get("valor_antes_iva"),
        valor_iva=phase1_dict.get("valor_iva"),
        valor_total=phase1_dict.get("valor_total"),
        forma_pago=phase1_dict.get("forma_pago"),
        plazo_entrega=phase1_dict.get("plazo_entrega"),
        garantia=phase1_dict.get("garantia"),
        anticipo=phase1_dict.get("anticipo"),
        pago_saldo=phase1_dict.get("pago_saldo"),
        items=[ItemCotizacion(**i) for i in (phase1_dict.get("items") or [])],
        nombre_archivo=nombre,
        campos_encontrados=phase1_dict.get("campos_encontrados", 0),
    )

    # Guardar archivo temporal (igual que antes)
    try:
        _COTIZACIONES_DIR.mkdir(parents=True, exist_ok=True)
        temp_path = _COTIZACIONES_DIR / f"temp_{solicitud_id}.{ext}"
        temp_path.write_bytes(contenido)
        resultado.temp_file_ext = ext
    except Exception as e:
        log.warning("[extraccion] No se pudo guardar archivo temporal %s: %s", solicitud_id, e)

    # ── Fase 2: background (Gemini) ───────────────────────────────────────────
    # Solo lanzar si hay campos vacíos que Gemini podría completar
    campos_vacios = sum(1 for v in [
        resultado.proveedor_nit, resultado.proveedor_nombre, resultado.valor_total,
        resultado.forma_pago, resultado.plazo_entrega, resultado.garantia,
    ] if v is None)

    if campos_vacios > 0 and settings.gemini_api_key:
        background_tasks.add_task(
            run_phase2, contenido, ext, str(solicitud_id), phase1_dict
        )
        log.info("[extraccion] Fase 2 programada para solicitud %s (%d campos vacíos)", solicitud_id, campos_vacios)

    return resultado


_COTIZACIONES_DIR = Path("/app/data/cotizaciones")


@router.get(
    "/solicitudes/{solicitud_id}/cotizacion/extraccion/resultado",
    status_code=status.HTTP_200_OK,
)
def get_extraccion_resultado(
    solicitud_id: uuid.UUID,
    current_user: User = Depends(require_compras),
):
    """Poll: retorna el resultado de Fase 2 (Gemini) cuando esté disponible.

    Retorna 204 si Fase 2 aún no terminó.
    Retorna 200 + JSON con campos completados cuando esté listo.
    """
    from app.services.extraction_pipeline import get_phase2_result
    from fastapi.responses import Response
    result = get_phase2_result(str(solicitud_id))
    if result is None:
        return Response(status_code=status.HTTP_204_NO_CONTENT)
    return result


@router.post(
    "/solicitudes/{solicitud_id}/cotizacion",
    response_model=CotizacionRead,
    status_code=status.HTTP_201_CREATED,
)
async def crear_cotizacion(
    solicitud_id: uuid.UUID,
    background_tasks: BackgroundTasks,
    payload: CotizacionCreate,
    current_user: User = Depends(require_compras),
    oc_db: Session = Depends(get_oc_db),
):
    solicitud = oc_db.get(SolicitudOC, solicitud_id)
    if not solicitud:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Solicitud no encontrada.")

    estados_validos = {EstadoOC.en_cotizacion, EstadoOC.rechazada}
    if solicitud.estado not in estados_validos:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail=f"No se puede cargar cotización en estado '{solicitud.estado}'.",
        )

    # Normalizar ítems: asignar número secuencial si no viene en el payload
    items_normalizados: Optional[list[dict]] = None
    if payload.items:
        items_normalizados = [
            {**item, "num": item.get("num") or (i + 1)}
            for i, item in enumerate(payload.items)
        ]

    cotizacion = CotizacionProveedor(
        solicitud_id=solicitud_id,
        proveedor_nombre=payload.proveedor_nombre,
        proveedor_nit=payload.proveedor_nit,
        proveedor_email=payload.proveedor_email,
        numero_cotizacion_proveedor=payload.numero_cotizacion_proveedor,
        valor_unitario=payload.valor_unitario,
        valor_antes_iva=payload.valor_antes_iva,
        valor_iva=payload.valor_iva,
        valor_total=payload.valor_total,
        fecha_estimada_entrega=payload.fecha_estimada_entrega,
        forma_pago=payload.forma_pago,
        plazo_entrega=payload.plazo_entrega,
        garantia=payload.garantia,
        anticipo=payload.anticipo,
        pago_saldo=payload.pago_saldo,
        observaciones=payload.observaciones,
        items=items_normalizados,
        extraccion_automatica=False,
        created_at=datetime.now(timezone.utc),
    )
    oc_db.add(cotizacion)

    # Avanzar estado de la solicitud a pendiente_aprobacion
    estado_anterior = solicitud.estado
    solicitud.estado = EstadoOC.pendiente_aprobacion
    solicitud.fecha_cotizacion = datetime.now(timezone.utc)
    solicitud.updated_at = datetime.now(timezone.utc)
    oc_db.add(solicitud)

    registrar_cambio_estado(
        oc_db,
        solicitud.id,
        estado_anterior,
        EstadoOC.pendiente_aprobacion,
        usuario_id=current_user.id,
        usuario_nombre=current_user.full_name,
    )

    oc_db.commit()
    oc_db.refresh(cotizacion)

    # Vincular el archivo temporal que dejó la extracción automática (si existe)
    # La extracción guarda: _COTIZACIONES_DIR/temp_{solicitud_id}.{ext}
    try:
        temp_candidates = list(_COTIZACIONES_DIR.glob(f"temp_{solicitud_id}.*"))
        if temp_candidates:
            temp_path = temp_candidates[0]
            ext = temp_path.suffix.lstrip(".")
            dest = _COTIZACIONES_DIR / f"{cotizacion.id}.{ext}"
            temp_path.rename(dest)
            cotizacion.pdf_path = str(dest)
            oc_db.add(cotizacion)
            oc_db.commit()
            oc_db.refresh(cotizacion)
    except Exception as e:
        log.warning("[cotizacion] No se pudo vincular archivo para cotización %s: %s", cotizacion.id, e)

    # Disparar emails Flujo 2 y 3 (cotización lista → pendiente_aprobacion)
    background_tasks.add_task(send_cotizacion_lista, solicitud)
    background_tasks.add_task(send_aprobacion_directora, solicitud, cotizacion)

    return cotizacion


@router.get(
    "/solicitudes/{solicitud_id}/cotizaciones",
    response_model=list[CotizacionRead],
)
def listar_cotizaciones(
    solicitud_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    oc_db: Session = Depends(get_oc_db),
):
    solicitud = oc_db.get(SolicitudOC, solicitud_id)
    if not solicitud:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Solicitud no encontrada.")

    cotizaciones = oc_db.exec(
        select(CotizacionProveedor)
        .where(CotizacionProveedor.solicitud_id == solicitud_id)
        .order_by(CotizacionProveedor.created_at.desc())
    ).all()
    return cotizaciones


@router.patch(
    "/cotizaciones/{cotizacion_id}/aprobar",
    response_model=CotizacionRead,
)
def aprobar_cotizacion(
    cotizacion_id: uuid.UUID,
    payload: AprobarPayload,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    oc_db: Session = Depends(get_oc_db),
):
    if not user_has_permission(db, current_user, "mod_oc_aprobar"):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Se requiere permiso para aprobar cotizaciones (mod_oc_aprobar).",
        )

    cotizacion = oc_db.get(CotizacionProveedor, cotizacion_id)
    if not cotizacion:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Cotización no encontrada.")

    cotizacion.aprobada = True
    cotizacion.valor_aprobado = payload.valor_aprobado
    cotizacion.aprobado_por_id = current_user.id
    cotizacion.observaciones_aprobacion = payload.observaciones_aprobacion
    oc_db.add(cotizacion)

    # Avanzar estado de la solicitud
    solicitud = oc_db.get(SolicitudOC, cotizacion.solicitud_id)
    if solicitud:
        estado_anterior = solicitud.estado
        solicitud.estado = EstadoOC.aprobada
        solicitud.fecha_aprobacion = datetime.now(timezone.utc)
        solicitud.aval_compra = current_user.full_name  # quien autorizó la compra
        solicitud.updated_at = datetime.now(timezone.utc)
        oc_db.add(solicitud)
        registrar_cambio_estado(
            oc_db,
            solicitud.id,
            estado_anterior,
            EstadoOC.aprobada,
            usuario_id=current_user.id,
            usuario_nombre=current_user.full_name,
        )

        # Si la solicitud requiere anticipo/proforma, notificar a Financiera
        if solicitud.tiene_proforma:
            background_tasks.add_task(
                send_proforma_financiero,
                solicitud,
                cotizacion,
                solicitud.proforma_path,
            )

    oc_db.commit()
    oc_db.refresh(cotizacion)
    return cotizacion


@router.patch(
    "/cotizaciones/{cotizacion_id}/rechazar",
    response_model=CotizacionRead,
)
def rechazar_cotizacion(
    cotizacion_id: uuid.UUID,
    payload: RechazarPayload,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    oc_db: Session = Depends(get_oc_db),
    db: Session = Depends(get_db),
):
    from app.services import email_service

    if not user_has_permission(db, current_user, "mod_oc_aprobar"):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Se requiere permiso para rechazar cotizaciones (mod_oc_aprobar).",
        )

    cotizacion = oc_db.get(CotizacionProveedor, cotizacion_id)
    if not cotizacion:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Cotización no encontrada.")

    cotizacion.aprobada = False
    cotizacion.aprobado_por_id = current_user.id
    cotizacion.observaciones_aprobacion = payload.observaciones_aprobacion
    oc_db.add(cotizacion)

    # Regresar solicitud a en_cotizacion para que el auxiliar busque otra cotización
    solicitud = oc_db.get(SolicitudOC, cotizacion.solicitud_id)
    if solicitud:
        estado_anterior = solicitud.estado
        solicitud.estado = EstadoOC.en_cotizacion
        solicitud.updated_at = datetime.now(timezone.utc)
        oc_db.add(solicitud)
        registrar_cambio_estado(
            oc_db,
            solicitud.id,
            estado_anterior,
            EstadoOC.en_cotizacion,
            usuario_id=current_user.id,
            usuario_nombre=current_user.full_name,
            notas=payload.observaciones_aprobacion,
        )

    oc_db.commit()
    oc_db.refresh(cotizacion)

    # Notificar al auxiliar por email
    if solicitud and solicitud.auxiliar_id:
        auxiliar = db.get(User, solicitud.auxiliar_id)
        if auxiliar and auxiliar.email:
            background_tasks.add_task(
                email_service.send_rechazo_cotizacion,
                solicitud,
                payload.observaciones_aprobacion,
                auxiliar.email,
            )
        else:
            log.warning(
                "[email] Rechazo cotización: auxiliar %s sin email en solicitud %s",
                solicitud.auxiliar_id,
                solicitud.consecutivo_os,
            )
    else:
        log.warning(
            "[email] Rechazo cotización: solicitud %s sin auxiliar_id asignado",
            solicitud.consecutivo_os if solicitud else cotizacion_id,
        )

    return cotizacion


@router.patch(
    "/cotizaciones/{cotizacion_id}/cancelar",
    response_model=CotizacionRead,
)
def cancelar_cotizacion(
    cotizacion_id: uuid.UUID,
    payload: CancelarCotizacionPayload,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    oc_db: Session = Depends(get_oc_db),
):
    """Cancela definitivamente la solicitud desde el panel de aprobación. KPI: rechazos_cotizacion."""
    if not user_has_permission(db, current_user, "mod_oc_aprobar"):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Se requiere permiso para cancelar cotizaciones (mod_oc_aprobar).",
        )

    cotizacion = oc_db.get(CotizacionProveedor, cotizacion_id)
    if not cotizacion:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Cotización no encontrada.")

    cotizacion.aprobada = False
    cotizacion.aprobado_por_id = current_user.id
    cotizacion.observaciones_aprobacion = payload.justificacion
    oc_db.add(cotizacion)

    solicitud = oc_db.get(SolicitudOC, cotizacion.solicitud_id)
    if solicitud:
        estado_anterior = solicitud.estado
        solicitud.estado = EstadoOC.cancelada
        solicitud.updated_at = datetime.now(timezone.utc)
        oc_db.add(solicitud)
        registrar_cambio_estado(
            oc_db,
            solicitud.id,
            estado_anterior,
            EstadoOC.cancelada,
            usuario_id=current_user.id,
            usuario_nombre=current_user.full_name,
            notas=payload.justificacion,
            tipo_accion="cancelacion_cotizacion",
        )

    oc_db.commit()
    oc_db.refresh(cotizacion)

    if solicitud:
        from app.services import email_service
        background_tasks.add_task(
            email_service.send_solicitud_rechazada,
            solicitud,
            payload.justificacion,
            current_user.full_name,
        )

    return cotizacion


@router.patch(
    "/cotizaciones/{cotizacion_id}/correccion",
    response_model=CotizacionRead,
)
def correccion_cotizacion(
    cotizacion_id: uuid.UUID,
    payload: CorreccionCotizacionPayload,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    oc_db: Session = Depends(get_oc_db),
    db: Session = Depends(get_db),
):
    """Manda a corrección: al auxiliar (nueva cotización) o al solicitante. KPI: reprocesos."""
    if not user_has_permission(db, current_user, "mod_oc_aprobar"):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Se requiere permiso para mandar a corrección (mod_oc_aprobar).",
        )

    if payload.destino not in ("auxiliar", "solicitante"):
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="destino debe ser 'auxiliar' o 'solicitante'.",
        )

    cotizacion = oc_db.get(CotizacionProveedor, cotizacion_id)
    if not cotizacion:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Cotización no encontrada.")

    cotizacion.aprobada = False
    cotizacion.aprobado_por_id = current_user.id
    cotizacion.observaciones_aprobacion = payload.que_corregir
    oc_db.add(cotizacion)

    solicitud = oc_db.get(SolicitudOC, cotizacion.solicitud_id)
    if solicitud:
        estado_anterior = solicitud.estado

        if payload.destino == "auxiliar":
            nuevo_estado = EstadoOC.en_cotizacion
            tipo_accion = "correccion_cotizacion_a_auxiliar"
        else:
            nuevo_estado = EstadoOC.en_correccion
            tipo_accion = "correccion_cotizacion_a_solicitante"

        solicitud.estado = nuevo_estado
        solicitud.updated_at = datetime.now(timezone.utc)
        # Cuando se manda al solicitante, almacenar la instrucción en la solicitud
        if payload.destino == "solicitante":
            solicitud.observaciones_compras = payload.que_corregir
        oc_db.add(solicitud)

        registrar_cambio_estado(
            oc_db,
            solicitud.id,
            estado_anterior,
            nuevo_estado,
            usuario_id=current_user.id,
            usuario_nombre=current_user.full_name,
            notas=payload.que_corregir,
            es_reproceso=True,
            tipo_accion=tipo_accion,
        )

    oc_db.commit()
    oc_db.refresh(cotizacion)

    # Notificar al auxiliar si el destino es auxiliar
    if solicitud and payload.destino == "auxiliar" and solicitud.auxiliar_id:
        from app.services import email_service
        auxiliar = db.get(User, solicitud.auxiliar_id)
        if auxiliar and auxiliar.email:
            background_tasks.add_task(
                email_service.send_rechazo_cotizacion,
                solicitud,
                payload.que_corregir,
                auxiliar.email,
            )

    return cotizacion


@router.patch(
    "/cotizaciones/{cotizacion_id}/editar",
    response_model=CotizacionRead,
)
def editar_cotizacion(
    cotizacion_id: uuid.UUID,
    payload: EditarCotizacionPayload,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    oc_db: Session = Depends(get_oc_db),
):
    """Permite al directivo/admin corregir datos de la cotización antes de aprobarla.
    Solo disponible mientras la cotización esté pendiente de revisión (aprobada == null).
    """
    if not user_has_permission(db, current_user, "mod_oc_aprobar"):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Se requiere permiso para editar cotizaciones (mod_oc_aprobar).",
        )

    cotizacion = oc_db.get(CotizacionProveedor, cotizacion_id)
    if not cotizacion:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Cotización no encontrada.")

    if cotizacion.aprobada is not None:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Solo se pueden editar cotizaciones pendientes de revisión.",
        )

    if payload.proveedor_nombre is not None:
        cotizacion.proveedor_nombre = payload.proveedor_nombre
    if payload.proveedor_nit is not None:
        cotizacion.proveedor_nit = payload.proveedor_nit
    if payload.proveedor_email is not None:
        cotizacion.proveedor_email = payload.proveedor_email
    if payload.numero_cotizacion_proveedor is not None:
        cotizacion.numero_cotizacion_proveedor = payload.numero_cotizacion_proveedor
    if payload.valor_antes_iva is not None:
        cotizacion.valor_antes_iva = payload.valor_antes_iva
    if payload.valor_iva is not None:
        cotizacion.valor_iva = payload.valor_iva
    if payload.valor_total is not None:
        cotizacion.valor_total = payload.valor_total
    if payload.forma_pago is not None:
        cotizacion.forma_pago = payload.forma_pago
    if payload.plazo_entrega is not None:
        cotizacion.plazo_entrega = payload.plazo_entrega
    if payload.observaciones is not None:
        cotizacion.observaciones = payload.observaciones
    if payload.items is not None:
        cotizacion.items = [
            {**item, "num": item.get("num") or (i + 1)}
            for i, item in enumerate(payload.items)
        ]

    oc_db.add(cotizacion)
    oc_db.commit()
    oc_db.refresh(cotizacion)
    return cotizacion


_ESTADOS_CORRECCION_DIRECTIVO = {
    EstadoOC.aprobada,
    EstadoOC.oc_enviada,
    EstadoOC.oc_en_plataforma,
    EstadoOC.entregada,
    EstadoOC.cerrada,
}

# Estados donde el proceso ya concluyó — requieren justificación extra al corregir
_ESTADOS_POST_CIERRE = {EstadoOC.entregada, EstadoOC.cerrada}


@router.patch(
    "/cotizaciones/{cotizacion_id}/corregir-directivo",
    response_model=CotizacionRead,
)
def corregir_directivo(
    cotizacion_id: uuid.UUID,
    payload: CorregirDirectivoPayload,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    oc_db: Session = Depends(get_oc_db),
):
    """Corrección directiva post-aprobación.

    Permite al director/admin corregir datos de la cotización ya aprobada
    (estados: aprobada, oc_enviada, oc_en_plataforma) sin cambiar el estado
    de la solicitud. Notifica al auxiliar y al solicitante.
    """
    if not user_has_permission(db, current_user, "mod_oc_aprobar"):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Se requiere permiso mod_oc_aprobar para realizar correcciones directivas.",
        )

    cotizacion = oc_db.get(CotizacionProveedor, cotizacion_id)
    if not cotizacion:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Cotización no encontrada.")

    solicitud = oc_db.get(SolicitudOC, cotizacion.solicitud_id)
    if not solicitud:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Solicitud no encontrada.")

    if solicitud.estado not in _ESTADOS_CORRECCION_DIRECTIVO:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=(
                f"La solicitud debe estar en uno de los estados: "
                f"{', '.join(e.value for e in _ESTADOS_CORRECCION_DIRECTIVO)}. "
                f"Estado actual: {solicitud.estado}."
            ),
        )

    es_post_cierre = solicitud.estado in _ESTADOS_POST_CIERRE
    if es_post_cierre and (not payload.motivo_post_cierre or len(payload.motivo_post_cierre.strip()) < 3):
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Se requiere una justificación para corregir un proceso ya cerrado.",
        )

    # Actualizar solo los campos provistos
    if payload.proveedor_nombre is not None:
        cotizacion.proveedor_nombre = payload.proveedor_nombre
    if payload.proveedor_nit is not None:
        cotizacion.proveedor_nit = payload.proveedor_nit
    if payload.proveedor_email is not None:
        cotizacion.proveedor_email = payload.proveedor_email
    if payload.numero_cotizacion_proveedor is not None:
        cotizacion.numero_cotizacion_proveedor = payload.numero_cotizacion_proveedor
    if payload.valor_antes_iva is not None:
        cotizacion.valor_antes_iva = payload.valor_antes_iva
    if payload.valor_iva is not None:
        cotizacion.valor_iva = payload.valor_iva
    if payload.valor_total is not None:
        cotizacion.valor_total = payload.valor_total
    if payload.forma_pago is not None:
        cotizacion.forma_pago = payload.forma_pago
    if payload.plazo_entrega is not None:
        cotizacion.plazo_entrega = payload.plazo_entrega
    if payload.observaciones is not None:
        cotizacion.observaciones = payload.observaciones
    if payload.items is not None:
        cotizacion.items = [
            {**item, "num": item.get("num") or (i + 1)}
            for i, item in enumerate(payload.items)
        ]
    if payload.valor_aprobado is not None:
        # Preservar el valor original la primera vez que se corrige
        if cotizacion.valor_aprobado_original is None and cotizacion.valor_aprobado != payload.valor_aprobado:
            cotizacion.valor_aprobado_original = cotizacion.valor_aprobado
        cotizacion.valor_aprobado = payload.valor_aprobado

    nota = (payload.observacion_correccion or "").strip() or None

    # Actualizar quién realizó la última modificación y su nota.
    cotizacion.aprobado_por_id = current_user.id
    if nota:
        cotizacion.observaciones_aprobacion = nota
    oc_db.add(cotizacion)

    # Registrar en historial sin cambiar estado
    estado_actual = solicitud.estado
    notas_historial = nota or "Corrección directiva sin observación."
    if es_post_cierre and payload.motivo_post_cierre:
        notas_historial = (
            f"[CORRECCIÓN POST-CIERRE] Justificación: {payload.motivo_post_cierre.strip()}\n"
            f"Observación: {nota or '(sin observación)'}"
        )
    registrar_cambio_estado(
        oc_db,
        solicitud.id,
        estado_actual,
        estado_actual,
        usuario_id=current_user.id,
        usuario_nombre=current_user.full_name,
        notas=notas_historial,
        es_reproceso=True,
        tipo_accion="correccion_directivo_post_cierre" if es_post_cierre else "correccion_directivo",
    )

    oc_db.commit()
    oc_db.refresh(cotizacion)

    # Notificar al auxiliar y al solicitante
    auxiliar_email: Optional[str] = None
    if solicitud.auxiliar_id:
        auxiliar = db.get(User, solicitud.auxiliar_id)
        if auxiliar and auxiliar.email:
            auxiliar_email = auxiliar.email

    if solicitud.estado in (EstadoOC.oc_enviada, EstadoOC.oc_en_plataforma):
        regen = regenerar_pdf_orden_por_solicitud(oc_db, db, solicitud.id)
        if regen:
            orden, cot_actualizada = regen
            email_prov = orden.email_proveedor or cot_actualizada.proveedor_email
            if orden.enviada_proveedor and email_prov:
                background_tasks.add_task(
                    send_oc_a_proveedor,
                    solicitud,
                    orden.numero_oc,
                    orden.pdf_path,
                    email_prov,
                    cot_actualizada.items,
                    auxiliar_email,
                )

    background_tasks.add_task(
        send_correccion_directivo,
        solicitud,
        cotizacion,
        nota or "",
        current_user.full_name,
        auxiliar_email,
    )

    # Alerta al propio directivo cuando corrige un proceso ya cerrado
    if es_post_cierre and current_user.email:
        background_tasks.add_task(
            send_alerta_correccion_post_cierre,
            solicitud,
            current_user.full_name,
            current_user.email,
            payload.motivo_post_cierre or "",
            nota or "",
        )

    return cotizacion


@router.get("/cotizaciones/{cotizacion_id}/pdf")
def descargar_pdf_cotizacion(
    cotizacion_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    oc_db: Session = Depends(get_oc_db),
):
    """Descarga el archivo PDF/Excel/Word adjunto a una cotización."""
    cotizacion = oc_db.get(CotizacionProveedor, cotizacion_id)
    if not cotizacion or not cotizacion.pdf_path:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Archivo no disponible.")
    path = Path(cotizacion.pdf_path)
    if not path.exists():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Archivo no encontrado en disco.")
    ext = path.suffix.lower().lstrip(".")
    _MIME = {
        "pdf": "application/pdf",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "xls": "application/vnd.ms-excel",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    media_type = _MIME.get(ext, "application/octet-stream")
    return FileResponse(str(path), media_type=media_type, filename=f"cotizacion_{cotizacion_id}.{ext}")
